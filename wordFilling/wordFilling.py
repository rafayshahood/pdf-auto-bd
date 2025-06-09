import json
from datetime import datetime, timedelta
from wordFilling.adjustDates import adjust_dates
from wordFilling.randomValGen import getRangeValuesArray
from wordFilling.docProcessing import process_document_full
import zipfile
import io
from docx import Document
import json
import random
import re

# Function to add 45 minutes to a time string
def add_45_minutes(time_str):
    time_obj = datetime.strptime(time_str, '%H:%M')
    new_time = time_obj + timedelta(minutes=45)
    return f"{time_obj.strftime('%H:%M')}-{new_time.strftime('%H:%M')}"

def remove_trailing_comma(lst):
    """
    Removes a trailing comma from the last element of a list if present.

    Parameters:
    - lst (list): A list of strings.

    Returns:
    - list: Updated list with no trailing comma in the last element.
    """
    if lst and lst[-1].endswith(","):
        lst[-1] = lst[-1].rstrip(",")  # Remove trailing comma
    return lst

def clean_safety_measures(safety_measures):
    """
    Removes 'COVID-19 Precaution' or 'COVID-19 Precautions' (case insensitive) from safety measures.
    Handles extra spaces or commas after removal.
    """
    return re.sub(r'covid-19 precaution[s]?,?\s*', '', safety_measures, flags=re.IGNORECASE).strip()

def remove_brackets(text):
    """
    Removes all square brackets [ ] from the given text.
    """
    return text.replace("[", "").replace("]", "")

def add_special_conditions(o2_flag, diabetec_flag):
    """
    Returns the additional condition-specific text if applicable.
    """
    condition_text = ""
    
    if o2_flag:
        condition_text += "Check O₂ saturation level with signs and symptoms of respiratory distress. "
    
    if diabetec_flag:
        condition_text += "SN to record blood sugar test results checked by Pt/PCG during the visits and report any significant changes to MD. SN to perform diabetic foot exam upon every visit. PCG assumes DM responsibilities, is confident, capable, and competent in checking blood sugar daily. "
    
    return condition_text

def modify_text2_with_conditions(text2, o2_flag, diabetec_flag):
    """
    Inserts O₂ and Diabetes-related sentences into text2 at the correct location.
    """
    target_lines = (
        "SN admitted the patient for comprehensive skilled nursing assessment, observation and evaluation of all body systems. "
        "SN to assess vital signs, pain level. SN performed to check vital signs and scale pain (1-10) every visit."
    )

    special_conditions_text = []
    if o2_flag:
        special_conditions_text.append("Check O₂ saturation level with signs and symptoms of respiratory distress.")
    if diabetec_flag:
        special_conditions_text.append(
            "SN to record blood sugar test results checked by Pt/PCG during the visits and report any significant changes to MD. "
            "SN to perform diabetic foot exam upon every visit. PCG assumes DM responsibilities, is confident, capable, and competent in checking blood sugar daily."
        )

    if special_conditions_text:
        conditions_text = " ".join(special_conditions_text)
        if target_lines in text2:
            text2 = text2.replace(target_lines, target_lines + " " + conditions_text, 1)
        else:
            text2 = conditions_text + " " + text2  # Fallback if target line is missing

    return text2


# # Function to ensure the dictionary has 9 pages
# def add_missing_pages(response_dict,  total_pages=9):
#     current_pages = list(response_dict.keys())
    
#     # Check how many pages are already in the dictionary
#     for i in range(1, total_pages + 1):
#         page_key = f"page{i}"
#         if page_key not in current_pages:
#             # Add missing pages with default values
#             response_dict[page_key] = {
#                 "text1": "Not enough disease to process",
#                 "text2": "Not enough disease to process",
#                 "med": "Not enough disease to process",
#                 "showButton": 1  
#             }
#     return response_dict


# --------------------------
# (Make sure that extractedResults, mainContResponse, and other external variables are defined.)
def fillDoc(submission_data, mainContResponse, selectedDiseaseList):
    # Example external variables (replace with your own logic or data)
    extractedResults = submission_data['extraction_results']
    constipation = extractedResults['diagnosis']['constipated']
    sn_name = submission_data['sn_name']
    appointment_dates = submission_data['appointment_dates']
    oxygenFlag = submission_data['extraction_results']['diagnosis']['oxygen']
    dm2_value =  submission_data['extraction_results']['diagnosis']['diabetec']
    # print(appointment_dates)
    appointment_times = submission_data['appointment_times']
    action = submission_data['action']




    nutritionalReq = extractedResults['extraDetails'].get('nutritionalReq', "")
    nutritionalReqCont = extractedResults['extraDetails'].get('nutritionalReqCont', "")

    # Only include `nutritionalReqCont` if it is not empty
    if nutritionalReqCont.strip():  
        nutritionalReqs = f"{nutritionalReq}, {nutritionalReqCont}"
    else:
        nutritionalReqs = nutritionalReq

    dischargeLastPage = {'text1': """Upon today’s assessment patient's condition is stable, vital signs remain stable recently. Patient/PCG monitored with discharge instruction.""",
                         'text2': """SN admitted the patient for comprehensive skilled nursing assessment, observation and evaluation of all body systems. SN to assess vital signs, pain level. SN performed to check vital signs and scale pain (1-10) every visit. """ + add_special_conditions(oxygenFlag, dm2_value) + """SN to evaluate therapeutic response to current/new medications and compliance to medication/diet regimen, home safety issues and psychosocial adjustment. SN informed Patient/PCG regarding possible discharge from services next visit. Patient/ PCG instructed re medication regimen -take all prescribed medications as ordered; if a dose is skipped never take double dose; do not stop taking medicine abruptly, keep your medicine in original container. Instructions are: measures to increase activity tolerance -use energy saving techniques, rest frequently during an activity, schedule an activity when most tolerated-after rest periods, after pain meds, at least one hour after meals; put most frequently used items within easy reach; eat a well-balanced diet; set realistic goals.""",
                         }
    print("a")


    current_pages = list(mainContResponse.keys())
    
    # Check how many pages are already in the dictionary
    for i in range(1, 9 + 1):
        page_key = f"page{i}"
        if page_key not in current_pages:
            # Add missing pages with default values
            mainContResponse[page_key] = {
                "text1": "Not enough disease to process",
                "text2": "Not enough disease to process",
                "med": "Not enough disease to process",
                "showButton": 1  
            }
            selectedDiseaseList.append("No disease")

            
    selectedDiseaseList.append("Discharge")

    print("b")

    # Modify text2 for all pages in mainContResponse
    # for page_key in mainContResponse:
    #     response_data = mainContResponse[page_key]
    #     response_data["text2"] = modify_text2_with_conditions(response_data["text2"], oxygenFlag, dm2_value)
    #     mainContResponse[page_key] = json.dumps(response_data)

    for page_key in list(mainContResponse):
        raw = mainContResponse[page_key]
        if isinstance(raw, str):
            # it’s a JSON string → parse it
            response_data = json.loads(raw)
        else:
            # it’s already a dict
            response_data = raw

        # now clean your text fields
        response_data["text1"] = remove_brackets(response_data["text1"])
        response_data["text2"] = remove_brackets(response_data["text2"])

        # save it back as a dict (or re-dump to JSON if you really need a string later)
        mainContResponse[page_key] = response_data
    
    # for page_key in mainContResponse:
    #     response_data =json.loads( mainContResponse[page_key]) # Parse JSON
    #     # Remove brackets from text1 and text2
    #     response_data["text1"] = remove_brackets(response_data["text1"])
    #     response_data["text2"] = remove_brackets(response_data["text2"])
        

    #     # Save back the cleaned data
    #     mainContResponse[page_key] = json.dumps(response_data)


    print("d")

    wordFileName = './wordFilling/1-page-version copy.docx'  # Ensure this file exists in your working directory


    # Get Action
    getAction = action
    if getAction == "Reset":
        valuesToGet = 9
        # wordFileName = './noteTemplates/9-page-version.docx'  # Ensure this file exists in your working directory
    elif getAction == "Discharge":
        valuesToGet = 10
        # wordFileName = './noteTemplates/10-page-version.docx'  # Ensure this file exists in your working directory
        mainContResponse['page10'] = json.dumps(dischargeLastPage)

    # Generate arrays for 10 iterations.
    tempArray = []
    hrArray = []
    rrArray = []
    oxygenArray = []
    bsFastArray = []
    bsRapidArray = []
    random_bpArray = []

    tempArray, hrArray, rrArray, oxygenArray, bsFastArray, bsRapidArray, random_bpArray = getRangeValuesArray(valuesToGet)

    if extractedResults['extraDetails']['can'] == "true" and extractedResults['extraDetails']['walker'] == "true":
        canWalkerText = 'cane, walker'
    elif extractedResults['extraDetails']['can'] == "true" and extractedResults['extraDetails']['walker'] == "false":
        canWalkerText = 'cane'
    elif extractedResults['extraDetails']['can'] == "false" and extractedResults['extraDetails']['walker'] == "false":
        canWalkerText = 'walker'
    else:
        canWalkerText = ''

    painScaleArray = ['5/10', '4/10', '4/10', '4/10', '4/10', '3/10', '3/10', '3/10']

    if getAction == "Discharge":
        painScaleArray.append('3/10')
        painScaleArray.append('2/10')
    if getAction == "Reset":
        painScaleArray.append('4/10')

    # Extract Patient Name for file naming
    patient_name = extractedResults['patientDetails']["name"] 
    zip_filename = f"{patient_name}.zip"  

    print("e")
    output_files = []

    for i in range(valuesToGet):
        headerPage = extractedResults['patientDetails']['providerName']
        check_f=False
        check_r=False

        print("e-a")

        # Extract the starting time from the appointment time string
        time_str = appointment_times[i]
        start_time_str = time_str.split('-')[0].strip() if '-' in time_str else time_str.strip()

        print("e-b")
        if getAction == "Reset" and i == 8:
            bsFastArray[i] = str(int(bsFastArray[i-1]) + random.randint(4, 8))
            bsRapidArray[i] = str(int(bsRapidArray[i-1]) + random.randint(4, 8))

    
        if datetime.strptime(start_time_str, "%H:%M") < datetime.strptime("10:00", "%H:%M"):
            bsValue = bsFastArray[i]
            check_f=True
        else:
            bsValue = bsRapidArray[i]
            check_r=True

        replacements_first_col = {
            'cane, walker': canWalkerText,
            'Pain in Lower back, Neck, Joints': extractedResults['diagnosis']['painIn'],
            '4/10': painScaleArray[i],
            'tpainmedhere': extractedResults['medications']['painMedications'],
            '05/07/23': adjust_dates(appointment_dates[i], appointment_times[i], constipation),
            'NAS, Controlled carbohydrate Low fat, Low cholesterol, NCS, Dash': nutritionalReqs
            }
        # Conditionally add the oxygen value only if oxygenFlag is True
        if oxygenFlag:
            replacements_first_col['OXVAL)('] = oxygenArray[i]+ "%"
        else:
            replacements_first_col['OXVAL)('] = ""
        print("e-b")

        # inside your for-i loop
        raw_page = mainContResponse.get(f'page{i+1}', {})
        if isinstance(raw_page, str):
            try:
                page_data = json.loads(raw_page)
            except json.JSONDecodeError:
                page_data = {}
        else:
            page_data = raw_page

        new_text1 = page_data.get("text1", "")
        replacement_text = page_data.get("text2", "")

        replacements_second_col = {
            'T- 96.8': "T- " + tempArray[i],
            'HR- 66': "HR- " + hrArray[i],
            'RR -16': "RR - " + rrArray[i],
            'Sitting 142/89': "Sitting " + random_bpArray[i],
            'PARKER, PETER/LVN': sn_name,
            'MR# 022-001': "MR# " + extractedResults['patientDetails']['medicalRecordNo'][-7:],
            'PORK, JOHN': extractedResults['patientDetails']["name"],
            '05/08/2023': datetime.strptime(appointment_dates[i], "%Y-%m-%d").strftime("%m/%d/%y"),
            '18:00-18:45': add_45_minutes(appointment_times[i]),
            'new text1': new_text1,
            'replacement text': replacement_text,

        }

        print("e-c")
        # bs value only shown if patient is diabetec
        if dm2_value:
            replacements_second_col['BS 198'] = "BS " + bsValue
        else:
            replacements_second_col['BS 198'] = "BS " 


        depressed_value = extractedResults['diagnosis']['depression']
        edemaResults = extractedResults['extraDetails']['edema']
        check_vertigo = extractedResults['extraDetails']['vertigo']
        palpitation_check = extractedResults['extraDetails']['palpitation']
        # remove covid 19-precautions
        allSafetyMeasures = extractedResults['extraDetails']['safetyMeasures']
        allSafetyMeasures = clean_safety_measures(allSafetyMeasures)

        print("f")
        out_file = process_document_full(
            wordFileName, headerPage, selectedDiseaseList[i], extractedResults, replacements_first_col, replacements_second_col,
            allSafetyMeasures, dm2_value, edemaResults, depressed_value, i, getAction, valuesToGet, check_vertigo, 
            check_f, check_r, palpitation_check
        )

        print("g")
        if out_file:  # Ensure out_file is valid
            output_files.append(out_file)
        else:
            print(f"Failed to process document for page {i}")

        print("h")
    return output_files
    #         # Create an in-memory ZIP file
    # zip_buffer = io.BytesIO()

    # with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
    #     for file in output_files:
    #         zipf.write(file, arcname=file)  # Add each file to the ZIP

    # # Move to the beginning of the buffer
    # zip_buffer.seek(0)

    # # Provide a single download button for the ZIP file
    # st.download_button(
    #     label="Download All Notes",
    #     data=zip_buffer,
    #     file_name=zip_filename,
    #     mime="application/zip"
    # )



    # def merge_word_documents(doc_files):
    #     """
    #     Merges multiple Word documents while preserving layout, styles, and formatting,
    #     ensuring that extra blank pages do not appear.

    #     Parameters:
    #     - doc_files (list): List of paths to .docx files to be merged.

    #     Returns:
    #     - io.BytesIO: In-memory file object of the merged document.
    #     """
    #     # Load the first document as the base
    #     merged_doc = Document(doc_files[0])

    #     for file in doc_files[1:]:
    #         doc = Document(file)

    #         # Ensure proper spacing and avoid unwanted blank pages
    #         if merged_doc.paragraphs[-1].text.strip():  # If last paragraph has text, add a page break
    #             merged_doc.add_page_break()

    #         # Append content while preserving structure
    #         for element in doc.element.body:
    #             merged_doc.element.body.append(element)

    #     # Save the merged document to a BytesIO buffer
    #     merged_buffer = io.BytesIO()
    #     merged_doc.save(merged_buffer)
    #     merged_buffer.seek(0)
        
    #     return merged_buffer

    # # Merge all generated documents
    # merged_buffer = merge_word_documents(output_files)

    # # Provide a download button for the merged document
    # st.download_button(
    #     label="Download Merged Document",
    #     data=merged_buffer,
    #     file_name="merged_notes.docx",
    #     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    # )
